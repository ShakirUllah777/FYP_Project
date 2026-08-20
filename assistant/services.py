"""
Shared AI service layer.

Every function here tries a real LLM call (Anthropic's Claude, via the
`anthropic` SDK) first, using the ANTHROPIC_API_KEY from settings/.env.
If no key is configured, or the API call fails for any reason (offline
dev machine, quota, network policy, etc.), each function falls back to a
deterministic, dependency-free implementation so the project always
keeps working end-to-end for a live demo or a viva - it just becomes
"smart" instead of "rule based" automatically once a key is present.
"""
import json
import difflib
from django.conf import settings

MODEL_NAME = 'claude-sonnet-4-6'


def _get_client():
    api_key = getattr(settings, 'ANTHROPIC_API_KEY', '') or ''
    if not api_key:
        return None
    try:
        import anthropic
        return anthropic.Anthropic(api_key=api_key)
    except Exception:
        return None


def _extract_text(response):
    parts = []
    for block in getattr(response, 'content', []):
        if getattr(block, 'type', None) == 'text':
            parts.append(block.text)
    return '\n'.join(parts).strip()


# ---------------------------------------------------------------------
# Rule-based fallbacks (identical to the project's original behaviour)
# ---------------------------------------------------------------------

FALLBACK_SUGGESTIONS = {
    'machine learning': {
        'title': 'AI-Powered Prediction System Using Machine Learning',
        'description': 'This project focuses on building an intelligent prediction system using machine learning algorithms. The system will analyse historical data, identify patterns, and provide accurate predictions to help users make better decisions.'
    },
    'web': {
        'title': 'Full Stack Web Application with Modern Technologies',
        'description': 'A complete web-based platform built using modern frontend and backend technologies. The application will include user authentication, real-time data handling, and a clean responsive interface.'
    },
    'mobile': {
        'title': 'Cross-Platform Mobile Application for Students',
        'description': 'A mobile application developed for university students to manage their daily academic activities. The app will support both Android and iOS platforms with an intuitive user interface.'
    },
    'database': {
        'title': 'Intelligent Database Management and Analytics System',
        'description': 'A robust database management system with built-in analytics and reporting features. The project will handle large datasets efficiently and provide visual insights through interactive dashboards.'
    },
    'chatbot': {
        'title': 'AI Chatbot for University Student Support',
        'description': 'An intelligent conversational chatbot designed to assist university students with their queries. The system uses natural language processing to understand questions and provide helpful, accurate responses.'
    },
    'ecommerce': {
        'title': 'E-Commerce Platform with Smart Recommendation Engine',
        'description': 'A fully functional online shopping platform with an AI-powered product recommendation engine. The system will personalise user experience based on browsing history and purchase patterns.'
    },
    'security': {
        'title': 'Network Security Monitoring and Threat Detection System',
        'description': 'A cybersecurity tool that continuously monitors network traffic to detect and alert on potential threats. The system uses pattern recognition to identify suspicious activities in real time.'
    },
    'image': {
        'title': 'Image Recognition and Classification System Using Deep Learning',
        'description': 'A deep learning based image recognition system capable of classifying objects with high accuracy. The project will use convolutional neural networks trained on large image datasets.'
    },
    'healthcare': {
        'title': 'Smart Healthcare Management System with Patient Analytics',
        'description': 'A digital healthcare platform that manages patient records, appointments, and medical history. The system includes analytics to help doctors identify health trends and improve patient care.'
    },
    'default': {
        'title': 'Innovative Software Solution for Real World Problem',
        'description': 'A well-structured software project that addresses a real world challenge using modern technologies. The system will be designed with scalability, performance, and user experience as core priorities.'
    },
}

FALLBACK_CHAT_RESPONSES = {
    # --- General / Welcome ---
    'what is collabspace': 'CollabSpace is a platform built for Pakistani university students to find partners for their Final Year Projects (FYP) and paid tasks. You can post your project, browse others, and message students directly.',
    'what is this website': 'This is CollabSpace, a student collaboration platform where you can find FYP partners, post paid tasks, and connect with skilled students at your university.',
    'what is this platform': 'CollabSpace connects university students across Pakistan for Final Year Projects and paid freelance-style tasks, with skills, messaging, team workspaces, and reviews all built in.',
    'how does it work': 'It is simple. Register, add your skills, post your FYP idea or browse others, message the right student, and collaborate. Everything happens right here on the platform.',
    'who is this for': 'CollabSpace is built for university students - whether you need FYP teammates, want to offer paid help on a task, or want to explore communities of students with similar interests.',
    'who made this': 'CollabSpace is a Final Year Project built to help Pakistani university students find collaborators for their own FYPs and paid tasks.',
    'is it free': 'Yes. CollabSpace is completely free for all university students. Just register and start collaborating.',
    'which universities': 'CollabSpace is open to students from all Pakistani universities - FAST, NUST, COMSATS, UET, IBA, LUMS and more.',
    'who can join': 'Any currently enrolled university student in Pakistan can join CollabSpace by registering with their details and verifying their email.',
    'is registration open to everyone': 'Yes, registration is open to all university students. You just need a valid email address and your academic details to sign up.',
    'can i browse without an account': 'Some pages are public, but you will need to register and log in to post projects, message students, save searches, join communities, or leave reviews.',
    'is there a mobile app': 'CollabSpace currently runs as a website that works in your phone or laptop browser. There is no separate mobile app at this time.',
    'is collabspace open source': 'CollabSpace is a Final Year Project built by a student team, not a public open-source project.',
    'hello': 'Hello! Welcome to CollabSpace. I am here to help you. Ask me anything about how to use the platform!',
    'hi': 'Hi there! I am the CollabSpace assistant. How can I help you today?',
    'hey': 'Hey! Good to see you here. What would you like help with on CollabSpace?',
    'assalam o alaikum': 'Walaikum Assalam! Welcome to CollabSpace. How can I help you today?',
    'salam': 'Walaikum Assalam! I am the CollabSpace assistant - ask me anything about the platform.',
    'good morning': 'Good morning! How can I help you with CollabSpace today?',
    'help': 'Sure! Here are some things I can help with:\n- How to register or login\n- How to post a project\n- How to find and message students\n- How to add your skills\n- What FYP and Paid Tasks mean\n- Communities, Team Workspaces, reviews, and reporting\n\nJust ask your question!',
    'what can you do': 'I can answer questions about registering and logging in, posting or browsing projects, adding skills, using the AI Resume Extractor, saving posts and searches, communities, team workspaces, the leaderboard, reviews, and reporting. Just ask!',
    'thank you': 'You are welcome! Feel free to ask if you need anything else.',
    'thanks': 'Happy to help! Good luck with your FYP!',
    'bye': 'Goodbye! Come back anytime you need help with CollabSpace.',

    # --- Registration & Login ---
    'how to register': 'Click the Register button at the top. Fill in your name, email, university details, program, and semester. After registering, verify your email, then login and add your skills from your profile page.',
    'how to create an account': 'Click Register at the top of the page, fill in your name, university email, program, and semester, then verify your email before logging in.',
    'how to login': 'Click the Login button at the top. Enter your email and password, then enter the 6-digit code sent to your email. You will be redirected to the Tasks feed after logging in.',
    'how to sign up': 'Click Register, fill in your details including your university email, program and semester, then verify your email and log in.',
    'forgot password': 'Currently password reset is not available. Please contact your administrator or create a new account.',
    'how to change password': 'Changing your password is not self-service yet. Please contact your platform administrator for help.',
    'how to change email': 'Changing your registered email is not currently self-service. Please contact your platform administrator to update it.',
    'i did not receive otp': 'If you did not receive your 6-digit login code, check your spam folder first. If it still has not arrived, use the Resend Verification Email option, or wait a few minutes and try logging in again.',
    'how to verify email': 'After registering, check your inbox for a verification email and click the link inside. If you do not see it, go to your profile and click "Resend Verification Email".',
    'how to resend verification email': 'Go to your Profile page. If your email is not yet verified, you will see a "Resend Verification Email" button - click it to get a new verification link.',
    'what is otp': 'OTP stands for One-Time Password. It is the 6-digit code emailed to you during login as a second security step after your password.',
    'forgot username': 'CollabSpace uses your registered email address to log in, not a separate username, so there is nothing to forget - just use the email you registered with.',
    'how to delete account': 'Account deletion is not self-service yet - please contact your platform administrator to request account removal.',
    'delete account': 'Account deletion is not self-service yet - please contact your platform administrator to request account removal.',
    'is my data safe': 'Your profile, messages, and posts are only visible to registered CollabSpace users, and moderators can act on any reported content to keep the platform safe.',
    'contact support': 'For anything the assistant cannot resolve, please reach out to your CollabSpace administrator or supervisor directly.',

    # --- Profile & Skills ---
    'how to edit profile': 'Go to your Profile page and click Edit Skills - from there you can update your bio, links, department, program, semester, and batch.',
    'how to view profile': 'Click on any student username in the Tasks feed to view their full profile - skills, posts, ratings, department, and contact info.',
    'how to add skills': 'Go to your Profile page and click Edit Skills. You will see skills grouped by category. Check the ones you know and set your proficiency level.',
    'how to remove a skill': 'Go to your Profile page, click Edit Skills, then uncheck any skill you no longer want listed and save your changes.',
    'what skills are available': 'We have skills in several categories: Programming Languages, Frameworks, AI/Data, Databases, DevOps/Cloud, Mobile, Web/UI/UX Design, Security/QA, and Others.',
    'programming languages skills': 'The Programming Languages category covers languages like Python, Java, JavaScript, TypeScript, C, C++, C#, Go, PHP, Ruby, Rust, Swift, Kotlin, and more.',
    'frameworks skills': 'The Frameworks category covers tools like Django, Flask, FastAPI, React.js, Vue.js, Angular, Next.js, Node.js, Express.js, Laravel, Spring Boot, and more.',
    'ai and data skills': 'The AI, Data Science and ML category covers Machine Learning, Deep Learning, TensorFlow, PyTorch, Data Science, Data Analytics, NLP, Computer Vision, and Generative AI / LLMs.',
    'database skills': 'The Databases category covers MySQL, PostgreSQL, MongoDB, SQLite, Redis, Firebase, Oracle DB, Microsoft SQL Server, and more.',
    'devops skills': 'The DevOps & Cloud category covers Docker, Kubernetes, AWS, Azure, Google Cloud Platform, CI/CD Pipelines, Terraform, and Linux Administration.',
    'mobile development skills': 'The Mobile Development category covers Flutter, React Native, Android Development, iOS Development, and Cross-Platform Mobile Dev.',
    'design skills': 'The Web & UI/UX Design category covers UI/UX Design, Figma, Adobe XD, Wireframing & Prototyping, Frontend Development, and Responsive Web Design.',
    'security skills': 'The Cybersecurity & QA category covers Cybersecurity, Ethical Hacking & Pen Testing, Network Security, Cryptography, and Software Testing / QA.',
    'skill proficiency': 'When you add skills manually or via the Resume Extractor, you can set (or the AI will guess) a proficiency level - Beginner, Intermediate, or Advanced - based on how often it appears and the experience context in your resume.',
    'what is department field': 'Department is your academic department, such as IT, CS, DS, SE, or EE - it shows on your profile and is used to filter the leaderboard.',
    'what is program field': 'Program is your degree program, such as BSSE, BSCS, BSDS, or BSIT - it appears on your profile card.',
    'what is semester field': 'Semester shows which semester (1 through 7) you are currently in, so teammates know where you are in your degree.',
    'what is batch field': 'Batch is your cohort label, like "Fall 2022" or "Spring 2023" - it helps other students find classmates from their own intake.',
    'what is looking for field': 'The "Looking For" field on your profile lets others know whether you want an FYP Partner, Paid Task work, or Both, helping match the right opportunities to you.',
    'what is availability field': 'Availability shows how many hours a week you can commit - under 5, 5 to 10, 10 to 20, or 20+ hours - so teammates know your bandwidth.',
    'is verified badge meaning': 'A blue checkmark badge next to a student name means their university email has been verified, which builds trust when connecting with them.',
    'id verification': 'Upload a photo of your student card from your profile settings. An admin will review it and mark your account as ID Verified, which builds trust with other students.',
    'how to upload student id': 'From your profile settings, upload a clear photo of your student card. An admin will review it and mark your profile as ID Verified once approved.',

    # --- AI Resume Extractor / Auto-Fill ---
    'how to upload resume': 'Go to More \u2192 AI Resume Extractor, then drag and drop your PDF, DOCX, or TXT resume into the upload box (or paste the text instead). Click Extract & Preview to see what was detected.',
    'resume extractor': 'The AI Resume Extractor reads your uploaded CV, detects matching platform skills, guesses your proficiency level, and can auto-fill your bio, GitHub, and LinkedIn on your profile - all reviewable before saving.',
    'auto fill profile': 'Upload your resume on the AI Resume Extractor page (under More). CollabSpace will suggest skills plus a bio, GitHub, and LinkedIn link pulled from your CV. Review the suggestions, then click "Apply to My Profile".',
    'what file types for resume': 'You can upload PDF, DOCX, or TXT resumes, or simply paste your resume text directly into the box - both work with the AI Resume Extractor.',
    'how accurate is resume extraction': 'The extractor only matches skills that are clearly mentioned in your resume text using whole-word matching, so short skill names like "C", "R", or "Go" will not be picked up from unrelated words.',
    'how to re-scan resume': 'On your Profile page, click "Re-Scan Resume" (shown once you have already uploaded one) to run the extractor again with an updated file.',
    'where is my uploaded resume': 'If you have uploaded a resume before, you will see a "View uploaded file" link on your Profile page under the Resume section.',
    'can i edit detected skills before saving': 'Yes. After extraction, every detected skill appears as a checkbox you can uncheck, and your bio, GitHub, and LinkedIn suggestions are shown in editable text fields before you click Apply to My Profile.',
    'does resume extractor fill bio automatically': 'Yes. Alongside your skills, the extractor pulls a short bio, and your GitHub and LinkedIn links if they appear in your resume, all shown for review before saving.',

    # --- Posting Projects ---
    'how to post': 'After logging in, click Add Post in the navbar. Fill in the title, type (FYP or Paid Task), description, required skills, and deadline. Then submit!',
    'how to add post': 'Go to Add Post from the navbar. Fill in your project details and click Create Post. Your post will appear in the Tasks feed for other students.',
    'how to edit a post': 'Go to My Posts, find the post you want to change, and click Edit to update its title, description, skills, or deadline.',
    'how to delete a post': 'Go to My Posts, find the post you want to remove, and click Delete. This cannot be undone.',
    'how to mark a task complete': 'Open your post from My Posts and click "Mark as Complete" once the work is finished - this also unlocks reviews for that post.',
    'what is fyp': 'FYP stands for Final Year Project. It is the major project every university student completes in their final year. CollabSpace helps you find the right teammates for it.',
    'what is paid task': 'A Paid Task is a project where you pay a skilled student to complete specific work, like building a website, making a mobile app, or analysing data.',
    'fyp vs paid task': 'An FYP post is for finding teammates for your Final Year Project, while a Paid Task post is for hiring a student to complete specific work for payment.',
    'can i post more than one project': 'Yes, there is no limit on how many FYP or Paid Task posts you can create.',
    'how to set a deadline': 'When creating or editing a post, use the Deadline field to pick the date by which you need the work done or teammates found.',
    'what is ai suggest': 'The AI Suggester on the Add Post page helps you write a professional title and description. Just type a few keywords like "machine learning healthcare" and click Suggest.',
    'how does ai suggest work': 'Type a few keywords about your project idea into the AI Suggester box on the Add Post page, and it generates a professional title and description you can use or edit.',
    'ai scope estimate': 'When you create a post, CollabSpace automatically estimates its complexity (Beginner, Intermediate, or Advanced) and a rough timeline based on your description - shown on the post detail page.',
    'what does complexity mean': 'The complexity estimate (Beginner, Intermediate, or Advanced) reflects how technically demanding your project sounds, based on your description, and helps students gauge the effort involved.',
    'similar projects': 'When you submit a new FYP post, CollabSpace checks it against existing ideas and warns you if something very similar already exists, so your project stays original.',
    'duplicate idea': 'If your FYP description closely matches an existing post, you will see a warning after submitting with the similar post name and a similarity percentage.',
    'share a post': 'Every post has a public, shareable link (from the post detail page) with a rich preview for WhatsApp or LinkedIn, so you can share opportunities outside CollabSpace too.',
    'how to share a post publicly': 'Open any post detail page and copy its public link - it works even for people without a CollabSpace account and shows a preview when shared on WhatsApp or LinkedIn.',

    # --- Task Feed / Browsing ---
    'how to browse tasks': 'Click Task Feed in the navbar to see all open FYP and Paid Task posts. Use the filters at the top to narrow down by type, skill, or sort order.',
    'how to filter tasks': 'On the Task Feed, use the "All Types" and "All Skills" dropdowns to filter posts, then click Filter to apply, or Clear to reset.',
    'what does most recent sort mean': 'The "Most Recent" sort order on the Task Feed shows the newest posts first.',
    'what does match sort mean': 'Sorting by Match ranks posts by how closely their required skills line up with the skills on your profile, showing your best fits first.',
    'match score': 'The percentage badge (e.g. "72% Match") on a post shows how well your added skills line up with the skills that post requires - the higher, the better the fit.',
    'how to clear filters': 'On the Task Feed, click the Clear button next to the filter dropdowns to reset back to all types and all skills.',
    'how to save a search': 'Apply filters on the Tasks Feed and click "Save this search". You can revisit it anytime from More \u2192 Saved Searches instead of re-typing filters.',
    'saved search': 'On the Tasks Feed, set your filters (type, skill, sort) then click "Save this search" to store it. Access all your saved searches from More \u2192 Saved Searches.',
    'how to delete a saved search': 'Go to More \u2192 Saved Searches, find the search you want to remove, and click the delete/remove option next to it.',

    # --- Saved Posts / Bookmarks ---
    'saved posts': 'Tap the bookmark icon on any post in the Tasks feed or post detail page to save it for later. Find everything you have saved under More \u2192 Saved Posts.',
    'how to save a post': 'Click the bookmark icon on a task card (or on the post detail page) to save it. You can view or remove saved posts anytime from More \u2192 Saved Posts.',
    'bookmark': 'Bookmarking lets you save any FYP idea or paid task to revisit later without losing it in the feed. Look for the bookmark icon on each post.',
    'how to remove a saved post': 'Click the bookmark icon again on a post you have already saved (it will look filled-in) to remove it from your Saved Posts.',
    'where are my saved posts': 'Go to More \u2192 Saved Posts in the navbar to see every post you have bookmarked.',

    # --- Messaging ---
    'how to message': 'Find a post you like in the Tasks feed and click the Message button on that post. You can also click on a student username to view their profile and message them from there.',
    'how to chat': 'Click Messages in the navbar to see your inbox. Click any contact to open the chat. Type your message and press Send.',
    'how to find messages': 'Click Messages in the top navbar to open your inbox. All your conversations will be listed there.',
    'how to search messages': 'Go to More \u2192 Search Messages, then type a keyword to find a specific past conversation quickly.',
    'search messages': 'The Search Messages tool (under More) lets you find any past conversation by keyword instead of scrolling through your whole inbox.',
    'can i message anyone': 'Yes, once logged in you can message any student on the platform, either from their profile page or from a post they have created.',

    # --- Communities ---
    'what is community': 'Communities are topic-based hubs, like Web Development, Cybersecurity, or Python Developers, where students with matching skills or interests can connect, see each other, and browse related open posts.',
    'how to join a community': 'Go to Community in the navbar, open a community that interests you, and click "Join Community". You will also auto-appear in communities matching skills already on your profile.',
    'how to leave a community': 'Open the community page and click the "Joined" button again to leave that community.',
    'what communities are available': 'CollabSpace currently has nine communities: Web Development, Cybersecurity, Python Developers, AI & Machine Learning, Mobile App Developers, Data Science & Analytics, DevOps & Cloud, UI/UX & Design, and Game Development.',
    'web development community': 'The Web Development community is for students working with HTML, CSS, JavaScript, React, Django, Flask, Node.js and similar web technologies.',
    'cybersecurity community': 'The Cybersecurity community is for students interested in ethical hacking, network security, cryptography, and penetration testing.',
    'python developers community': 'The Python Developers community brings together students skilled in Python, Django, Flask, FastAPI, and data tools like Pandas.',
    'ai and machine learning community': 'The AI & Machine Learning community is for students working with Machine Learning, Deep Learning, TensorFlow, PyTorch, and Generative AI / LLMs.',
    'mobile app developers community': 'The Mobile App Developers community is for students building with Flutter, React Native, Android, or iOS development.',
    'data science community': 'The Data Science & Analytics community is for students working with data analytics, data visualisation, and tools like R, Tableau, and Power BI.',
    'devops community': 'The DevOps & Cloud community is for students working with Docker, Kubernetes, AWS, Azure, GCP, and CI/CD pipelines.',
    'ui ux community': 'The UI/UX & Design community is for students focused on product design, Figma, Adobe XD, and wireframing.',
    'game development community': 'The Game Development community is for students building games or interactive experiences with Unity or Unreal Engine.',
    'how do i become a community member': 'You automatically qualify for a community once you add a matching skill to your profile, or you can manually click "Join Community" on that community page anytime.',
    'community open posts': 'Each community page shows a list of open FYP and Paid Task posts that require skills related to that community, so you can find relevant opportunities faster.',

    # --- Recommended Teammates & Leaderboard ---
    'recommended teammates': 'Go to More \u2192 Recommended Teammates to see students whose skills best complement yours for FYP or paid work, ranked by match score.',
    'how are teammates recommended': 'CollabSpace compares your skills to other students, prioritising people who bring new, complementary skills while also sharing some common ground with you.',
    'no teammate recommendations': 'If you have not added skills yet, or there are no strong matches, Recommended Teammates falls back to showing the platform\u2019s most active and highly-rated students instead.',
    'leaderboard': 'The Leaderboard (under More) ranks students by activity, ratings, and completed collaborations, so you can see who is most active and trusted on the platform.',
    'how is leaderboard score calculated': 'Your Leaderboard Activity Score combines posts created, completed collaborations, team memberships, and your average review rating and review count, each weighted differently.',
    'what is activity score': 'Activity Score is a single number on the Leaderboard reflecting how active and trusted you are - it rewards completed work and good ratings more than just post volume.',
    'how to filter leaderboard by department': 'On the Leaderboard page, use the department dropdown at the top right to filter the rankings to a specific department.',
    'what do gold silver bronze mean': 'On the Leaderboard, the top three students are shown on a podium with gold, silver, and bronze medals for ranks 1, 2, and 3.',

    # --- Team Workspace ---
    'what is a team workspace': 'Once you agree to collaborate on a post, click "Start Team Workspace" to get a shared task board, milestones, file sharing, and meeting scheduling for that project.',
    'how to start a team workspace': 'Open a post you are collaborating on and click "Start Team Workspace" to create a shared space with tasks, milestones, files, and meetings for that project.',
    'how to add a team member': 'Inside a Team Workspace, use the "Add Member" option to invite another student into that project team.',
    'team roles': 'A Team Workspace has two roles - Owner, who created the team, and Collaborator, for everyone else who joins.',
    'milestones': 'Inside a Team Workspace, click "Add Milestone" to set project checkpoints with target dates, so your team can track progress toward the FYP deadline.',
    'how to update a milestone': 'Open the Team Workspace, find the milestone in the list, and update its status or date as your project progresses.',
    'team tasks': 'Inside a Team Workspace you can add tasks, assign them, and update their status (To Do, In Progress, or Done) so everyone knows what is left to finish.',
    'how to add a task in team workspace': 'Open your Team Workspace and click "Add Task", then fill in the title and priority to add it to the shared board.',
    'task priority': 'When adding a task in a Team Workspace, you can set its priority as Low, Medium, or High to help the team focus on what matters most.',
    'schedule meeting': 'Open your Team Workspace and click "Schedule Meeting" to set a date, time, and agenda that all team members can see.',
    'upload files': 'In a Team Workspace, use the "Upload File" option to share documents, code, or reports with your teammates in one shared place.',
    'how to delete a file in team workspace': 'Inside the Team Workspace file list, click the delete option next to a file you want to remove from the shared space.',
    'can i be in more than one team': 'Yes, you can be part of multiple Team Workspaces at the same time, one for each post you are collaborating on.',

    # --- Reviews & Reports ---
    'how to leave a review': 'After collaborating with someone, visit their profile and click "Leave a Review" to rate them 1 to 5 stars with a comment.',
    'when can i leave a review': 'You can leave a review for someone once the post you collaborated on together has been marked as complete.',
    'can i review the same person twice': 'You can leave one review per person per post - once submitted for that project, you cannot submit another review for the same collaboration.',
    'how to report': 'On any profile, click "Report" to flag inappropriate behaviour to our moderation team.',
    'what reasons can i report for': 'You can report for spam or misleading content, harassment or abusive behaviour, scam or non-payment, a fake profile, inappropriate content, or another reason you specify.',
    'what happens after i report someone': 'Your report goes to the Moderation Queue, where staff review it and mark it as Under Review, Resolved, or Dismissed.',

    # --- Supervisor & Moderation ---
    'supervisor dashboard': 'Faculty/supervisor accounts can open More \u2192 Supervisor Dashboard to review and endorse student FYP posts across the platform.',
    'endorsement': 'A supervisor can endorse an FYP post from the Supervisor Dashboard - an endorsed idea signals it has been reviewed and approved by faculty.',
    'how to become a supervisor': 'Supervisor accounts are set up by an administrator for faculty members - please contact your platform administrator if you need supervisor access.',
    'moderation queue': 'Staff accounts can access More \u2192 Moderation Queue to review reported users or posts and take action to keep the community safe.',

    # --- General closing ---
    'is collabspace safe for students': 'CollabSpace requires email verification, offers optional ID verification, a review system, and a reporting/moderation process, all designed to keep interactions on the platform trustworthy.',
    'what happens after graduation': 'Your CollabSpace account and history stay accessible after graduation, though the platform is mainly designed around active FYP and student paid-task collaboration.',
    'can teachers use collabspace': 'Yes, faculty members can be set up with a Supervisor account to review and endorse student FYP posts through the Supervisor Dashboard.',
}


def _fallback_suggestion(keywords):
    keywords = (keywords or '').lower()
    for key, value in FALLBACK_SUGGESTIONS.items():
        if key in keywords:
            return value
    return FALLBACK_SUGGESTIONS['default']


def _fallback_chat(question):
    question = (question or '').lower().strip()
    for key, value in FALLBACK_CHAT_RESPONSES.items():
        if key in question:
            return value
    return ("I am not sure about that yet. Here are some things I can help with: registering, "
            "adding skills, the AI Resume Extractor & profile auto-fill, saving posts, saved searches, "
            "posting a project, messaging students, team workspaces, reviews, or FYP vs Paid Tasks. "
            "Try asking one of those!")


# ---------------------------------------------------------------------
# Public functions used by views
# ---------------------------------------------------------------------

def generate_post_suggestion(keywords):
    """Returns {'title': ..., 'description': ...} for the Add Post AI Suggester."""
    client = _get_client()
    if client is None:
        return _fallback_suggestion(keywords)

    try:
        prompt = (
            "You help university students write a professional Final Year Project or "
            "paid-task post for a student collaboration platform called CollabSpace. "
            f"Given these keywords: \"{keywords}\", respond with ONLY a JSON object "
            'with exactly two keys: "title" (max 12 words) and "description" '
            "(2-3 sentences, professional tone). No markdown, no extra text."
        )
        response = client.messages.create(
            model=MODEL_NAME,
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        text = _extract_text(response).strip()
        text = text.removeprefix('```json').removeprefix('```').removesuffix('```').strip()
        data = json.loads(text)
        if 'title' in data and 'description' in data:
            return {'title': data['title'], 'description': data['description']}
    except Exception:
        pass
    return _fallback_suggestion(keywords)


def chatbot_reply(question):
    """Returns a plain-text reply for the floating assistant widget."""
    client = _get_client()
    if client is None:
        return _fallback_chat(question)

    try:
        system_prompt = (
            "You are the CollabSpace assistant, a friendly helper for a Django platform "
            "that connects Pakistani university students for Final Year Projects (FYP) and "
            "paid tasks. Answer briefly (2-4 sentences), in plain text (no markdown), and "
            "stay strictly on topic: registering, login/OTP, skills, the AI Resume Extractor "
            "(upload a PDF/DOCX/TXT resume to auto-detect skills and auto-fill bio/GitHub/"
            "LinkedIn), saved posts (bookmarks), saved searches, recommended teammates, the "
            "leaderboard, posts, messaging, team workspaces, supervisor endorsements, "
            "moderation/reporting, reviews, and general platform usage. If asked something "
            "unrelated, politely redirect to what you can help with."
        )
        response = client.messages.create(
            model=MODEL_NAME,
            max_tokens=200,
            system=system_prompt,
            messages=[{"role": "user", "content": question}],
        )
        text = _extract_text(response)
        if text:
            return text
    except Exception:
        pass
    return _fallback_chat(question)


def estimate_complexity(description):
    """Heuristic project-scope estimator. Deterministic by design (no
    network dependency) so it always works instantly for every post,
    with an optional LLM refinement layered on top when a key exists."""
    description = description or ''
    word_count = len(description.split())

    advanced_keywords = ['deep learning', 'neural network', 'distributed', 'blockchain',
                          'real-time', 'real time', 'microservice', 'computer vision',
                          'reinforcement learning', 'nlp', 'kubernetes', 'scalable']
    beginner_keywords = ['simple', 'basic', 'static', 'landing page', 'crud', 'todo', 'to-do']

    text_lower = description.lower()
    advanced_hits = sum(1 for kw in advanced_keywords if kw in text_lower)
    beginner_hits = sum(1 for kw in beginner_keywords if kw in text_lower)

    if advanced_hits >= 1 or word_count > 120:
        complexity, timeline = 'advanced', 'Full semester (14-16 weeks)'
    elif beginner_hits >= 1 or word_count < 30:
        complexity, timeline = 'beginner', '2-3 weeks'
    else:
        complexity, timeline = 'intermediate', '4-8 weeks'

    client = _get_client()
    if client is not None:
        try:
            prompt = (
                f"A university project is described as:\n\"{description}\"\n\n"
                'Respond with ONLY a JSON object: {"complexity": "beginner|intermediate|advanced", '
                '"timeline": "short human readable estimate"}. No extra text.'
            )
            response = client.messages.create(
                model=MODEL_NAME, max_tokens=100,
                messages=[{"role": "user", "content": prompt}],
            )
            text = _extract_text(response).strip()
            text = text.removeprefix('```json').removeprefix('```').removesuffix('```').strip()
            data = json.loads(text)
            if data.get('complexity') in ('beginner', 'intermediate', 'advanced'):
                complexity = data['complexity']
                timeline = data.get('timeline', timeline)
        except Exception:
            pass

    return complexity, timeline


def find_similar_posts(new_title, new_description, queryset, threshold=0.55, limit=3):
    """Duplicate/plagiarism-style idea checker using difflib's sequence
    matcher over title+description text. No embeddings/API required, so
    it works instantly and offline for every single post created."""
    candidate_text = f"{new_title}\n{new_description}".lower()
    scored = []
    for post in queryset:
        existing_text = f"{post.title}\n{post.description}".lower()
        ratio = difflib.SequenceMatcher(None, candidate_text, existing_text).ratio()
        if ratio >= threshold:
            scored.append((post, round(ratio * 100, 1)))
    scored.sort(key=lambda t: t[1], reverse=True)
    return scored[:limit]


def _skill_word_boundary_pattern(skill_name):
    """Builds a regex that matches a skill name as a whole token, not as a
    substring — e.g. the skill 'C' should match "I know C." but NOT
    "Certificate" or "Communication"; 'Go' should match "I use Go" but not
    "Google" or "ongoing"."""
    import re
    escaped = re.escape(skill_name.strip())
    return re.compile(r'(?<![A-Za-z0-9])' + escaped + r'(?![A-Za-z0-9])', re.IGNORECASE)


def extract_skills_from_text(resume_text, known_skill_names):
    """Extracts a list of skill names mentioned in pasted CV/resume text.
    Tries the LLM first; falls back to whole-word matching against the
    platform's existing Skill list (never plain substring matching, which
    used to cause false positives like 'C', 'R' or 'Go' matching inside
    unrelated words such as "Certificate" or "Google")."""
    client = _get_client()
    if client is not None:
        try:
            prompt = (
                "Extract technical skills mentioned in this resume text. "
                f"Only choose from this exact list of known skills: {json.dumps(known_skill_names)}.\n\n"
                f"Resume text:\n{resume_text[:4000]}\n\n"
                'Respond with ONLY a JSON array of matching skill names from the list, e.g. ["Python", "Django"]. '
                'Only include a skill if it is genuinely and clearly mentioned - do not guess or infer skills '
                'that are not explicitly present in the text.'
            )
            response = client.messages.create(
                model=MODEL_NAME, max_tokens=300,
                messages=[{"role": "user", "content": prompt}],
            )
            text = _extract_text(response).strip()
            text = text.removeprefix('```json').removeprefix('```').removesuffix('```').strip()
            data = json.loads(text)
            if isinstance(data, list):
                return [s for s in data if s in known_skill_names]
        except Exception:
            pass

    # Fallback: whole-word / whole-token matching, case-insensitive.
    text_lower = (resume_text or '')
    matched = []
    for name in known_skill_names:
        if _skill_word_boundary_pattern(name).search(text_lower):
            matched.append(name)
    return matched


# ---------------------------------------------------------------------
# Resume / CV parsing (file upload -> plain text)
# ---------------------------------------------------------------------

def extract_text_from_resume_file(uploaded_file):
    """Turns an uploaded .pdf / .docx / .txt resume file into plain text.
    Returns '' (and never raises) on any parsing failure so the caller can
    show a friendly message instead of a 500 error.

    Uses a couple of defensive tricks that matter a lot for real-world
    resumes (as opposed to plain-text test files):
      - PDFs: pypdf alone fails to extract *any* text from a surprising
        number of real resumes (custom/subsetted fonts from Canva, Word's
        "Save as PDF", etc). If pypdf comes back empty, we retry with
        pdfminer.six, which uses a different, more tolerant text-extraction
        pipeline and recovers text pypdf misses.
      - DOCX: many resume templates lay out contact info / skills in a
        table (e.g. a sidebar), and `document.paragraphs` only walks
        top-level paragraphs - text inside tables is silently skipped.
        We additionally walk every table cell.
      - The uploaded file's stream position is always reset (seek(0))
        before each parsing attempt, since a failed first attempt can
        leave the pointer at EOF and make a retry silently read nothing.
    """
    import io
    import logging
    logger = logging.getLogger(__name__)
    name = (getattr(uploaded_file, 'name', '') or '').lower()

    # Read the whole upload into memory ONCE and work off plain io.BytesIO
    # copies from here on. This sidesteps stream-position bugs (a failed
    # first parse attempt can leave the original file pointer at EOF) and
    # also sidesteps library-compatibility quirks - some parsers (e.g.
    # pdfminer.six) require a genuine io.IOBase instance and reject
    # Django's UploadedFile wrapper outright.
    try:
        uploaded_file.seek(0)
    except Exception:
        pass
    try:
        raw_bytes = uploaded_file.read()
    except Exception:
        logger.warning('Failed to read uploaded resume %r', name, exc_info=True)
        return ''
    if not raw_bytes:
        return ''

    if name.endswith('.pdf'):
        text = ''
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw_bytes))
            text = '\n'.join((page.extract_text() or '') for page in reader.pages).strip()
        except Exception:
            logger.warning('pypdf failed to parse uploaded resume %r', name, exc_info=True)
            text = ''

        if not text:
            # Fallback engine: handles PDFs whose fonts/encoding pypdf
            # can't decode (very common with Canva/Word-exported resumes).
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract_text
                text = (pdfminer_extract_text(io.BytesIO(raw_bytes)) or '').strip()
            except Exception:
                logger.warning('pdfminer.six also failed to parse uploaded resume %r', name, exc_info=True)
                text = ''

        return text

    if name.endswith('.docx'):
        try:
            import docx
            document = docx.Document(io.BytesIO(raw_bytes))
            parts = [p.text for p in document.paragraphs]
            # Walk tables too - a lot of resume templates put the sidebar
            # (skills, contact info) inside a table, which `.paragraphs`
            # above skips entirely.
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            parts.append(p.text)
            return '\n'.join(part for part in parts if part).strip()
        except Exception:
            logger.warning('python-docx failed to parse uploaded resume %r', name, exc_info=True)
            return ''

    if name.endswith('.doc'):
        # Legacy binary Word format - python-docx can't read it, and there's
        # no lightweight pure-Python parser for it. Fail clearly rather than
        # falling through to the raw-bytes decode below (which would just
        # produce binary garbage instead of a helpful error).
        return ''

    # .txt or anything else readable as plain text
    for encoding in ('utf-8', 'utf-16', 'latin-1'):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode('utf-8', errors='ignore')


def guess_skill_proficiency(resume_text, skill_name):
    """Very lightweight heuristic: the more a skill is mentioned / the more
    it appears alongside experience-signal words, the higher the guessed
    proficiency. Deterministic, no network call needed. Uses whole-word
    matching so short skill names (e.g. 'C', 'R', 'Go') aren't over-counted
    from unrelated words."""
    text_lower = (resume_text or '').lower()
    mentions = len(_skill_word_boundary_pattern(skill_name).findall(resume_text or ''))

    strong_signals = ['years of experience', 'expert', 'advanced', 'led', 'built and deployed',
                       'production', 'senior']
    context_hits = sum(1 for kw in strong_signals if kw in text_lower)

    if mentions >= 3 or context_hits >= 2:
        return 'advanced'
    if mentions == 2 or context_hits == 1:
        return 'intermediate'
    return 'beginner'


def extract_profile_details_from_resume(resume_text):
    """Pulls a short bio + GitHub/LinkedIn links out of resume text so the
    profile can be auto-filled alongside the detected skills. Tries the LLM
    for a punchy one-line bio, and always falls back to regex-based
    extraction for the links (works even with no API key)."""
    import re

    github_match = re.search(r'(https?://)?(www\.)?github\.com/[A-Za-z0-9_.\-]+', resume_text or '', re.I)
    linkedin_match = re.search(r'(https?://)?(www\.)?linkedin\.com/in/[A-Za-z0-9_.\-]+', resume_text or '', re.I)

    def _normalize(url):
        if not url:
            return ''
        url = url.strip().rstrip('/.,')
        if not url.startswith('http'):
            url = 'https://' + url
        return url

    github_url = _normalize(github_match.group(0)) if github_match else ''
    linkedin_url = _normalize(linkedin_match.group(0)) if linkedin_match else ''

    bio = ''
    client = _get_client()
    if client is not None and resume_text:
        try:
            prompt = (
                "Write ONE short, professional bio sentence (max 20 words, no quotes, "
                "no markdown) summarising this student's resume for a student "
                f"collaboration platform profile:\n\n{resume_text[:3000]}"
            )
            response = client.messages.create(
                model=MODEL_NAME, max_tokens=80,
                messages=[{"role": "user", "content": prompt}],
            )
            bio = _extract_text(response).strip().strip('"')[:200]
        except Exception:
            bio = ''

    if not bio:
        # Fallback: use the first meaningful line of the resume as a bio seed.
        skip_prefixes = ('email', 'phone', 'address', 'github', 'linkedin', 'website', 'skills', 'contact')
        for raw_line in (resume_text or '').splitlines():
            line = raw_line.strip()
            if not line or 'http' in line.lower():
                continue
            candidate = line
            if ':' in line[:20]:
                prefix, _, remainder = line.partition(':')
                if prefix.strip().lower() in skip_prefixes:
                    continue
                candidate = remainder.strip() or line
            if len(candidate) > 25:
                bio = candidate[:200]
                break

    return {'bio': bio, 'github': github_url, 'linkedin': linkedin_url}
